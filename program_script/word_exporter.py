import requests, json
from docx import Document
from docx.shared import RGBColor

class WordExporter(object):
    def __init__(self, transcript_id, API_key):
        self.API_response = self.make_API_call(transcript_id, API_key)
        self.word_file = Document()
        self.error_code = None
        self.process_response(transcript_id)

    def make_API_call(self, transcript_id, API_key):
        url = 'https://api.capio.ai/v1/speech/transcript/' + transcript_id
        headers = {'apiKey': API_key}
        return requests.get(url, headers=headers)

    def process_response(self, transcript_id):
        if self.API_response.status_code == 200:
            self.API_response = self.API_response.json()
            self.write_transcript_to_file(10)
            self.word_file.save(transcript_id + '.docx')
        else:
            self.error_code = str(self.API_response.status_code)
            return self.error_code

    def write_transcript_to_file(self, space_after_timestamp):
        for sentence in self.API_response:
            line = self.word_file.add_paragraph()

            raw_time = sentence['result'][0]['alternative'][0]['words'][0]['from']
            time = self.reformat_time(raw_time)
            time_runner = line.add_run(time + " " * space_after_timestamp)
            time_runner.bold = True
            font = time_runner.font
            font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

            for word in sentence['result'][0]['alternative'][0]['words']:
                word_runner = line.add_run(word['word'] + " ")
                if word['confidence'] < 0.75:
                    font = word_runner.font
                    font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        return True

    def reformat_time(self, raw_time):
        raw_time = float(format(raw_time, '.2f'))
        hours = int(raw_time / 3600)
        raw_time -= hours * 3600
        minutes = int(raw_time / 60)
        raw_time -= minutes * 60
        seconds = int(raw_time)
        raw_time = format(raw_time - seconds, '.2f')
        centiseconds = int(str(raw_time)[2:])

        if hours < 10:
            formatted_hours = '0' + str(hours)
        else:
            formatted_hours = str(hours)
        if minutes < 10:
            formatted_minutes = '0' + str(minutes)
        else:
            formatted_minutes = str(minutes)
        if seconds < 10:
            formatted_seconds = '0' + str(seconds)
        else:
            formatted_seconds = str(seconds)
        if centiseconds < 10:
            formatted_centiseconds = '0' + str(centiseconds)
        else:
            formatted_centiseconds = str(centiseconds)

        return formatted_hours + ":" + formatted_minutes + ":" + formatted_seconds \
        + '.' + formatted_centiseconds
