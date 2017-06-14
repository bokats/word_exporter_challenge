from word_exporter import WordExporter
from docx import Document

class WordExporterTests(object):

    def test_correct_API_call(self):
        w = WordExporter('593f237fbcae700012ba8fcd', '262ac9a0c9ba4d179aad4c0b9b02120a')
        if w.error_code is not None:
            print('Fail correct API call test')

    def test_incorrect_API_key(self):
        w = WordExporter('593f237fbcae700012ba8fcd', '262ac9a0c02120a')
        if w.error_code is None:
            print('Fail incorrect API key test')

    def test_incorrect_transcript_id(self):
        w = WordExporter('593', '262ac9a0c9ba4d179aad4c0b9b02120a')
        if we.error_code is None:
            print('Fail incorrect transcript id test')

    def test_time_reformatting(self):
        w = WordExporter('593', '262ac9a0c9ba4d179aad4c0b9b02120a')
        if w.reformat_time(0.12421412) != '00:00:00.12':
            print('Time reformatting test fail', 0.12421412)
        if w.reformat_time(0.9988) != '00:00:01.00':
            print('Time reformatting test fail', 0.9988)
        if w.reformat_time(59.998) != '00:01:00.00':
            print('Time reformatting test fail', 59.998)
        if w.reformat_time(61.02) != '00:01:01.02':
            print('Time reformatting test fail', 61.02)
        if w.reformat_time(122.998) != '00:02:03.00':
            print('Time reformatting test fail', 122.598)
        if w.reformat_time(200) != '00:03:20.00':
            print('Time reformatting test fail', 200)
        if w.reformat_time(3599.998) != '01:00:00.00':
            print('Time reformatting test fail', 3599.998)

    def check_data_integrity(self):
        w = WordExporter('593f237fbcae700012ba8fcd', '262ac9a0c9ba4d179aad4c0b9b02120a')
        doc = Document('593f237fbcae700012ba8fcd.docx')

        if not doc:
            print('No document created test fail')
            return

        if len(w.API_response) != len(doc.paragraphs):
            print('Number of lines mismatch test fail')
            return

        index = 0
        for paragraph in doc.paragraphs:
            paragraph = paragraph.text.split(" " * 10)
            timestamp = w.API_response[index]['result'][0]['alternative'][0]['words'][0]['from']
            if w.reformat_time(timestamp) != paragraph[0]:
                print('Time data mismatch', timestamp)

            sentence = w.API_response[index]['result'][0]['alternative'][0]['transcript']
            if sentence != paragraph[1][:-1]:
                print('Sentence data mismatch', sentence)
            index += 1

    def run_tests(self):
        self.test_correct_API_call()
        self.test_incorrect_API_key()
        self.test_time_reformatting()
        self.check_data_integrity()

t = WordExporterTests()
t.run_tests()
